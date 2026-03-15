#!/usr/bin/env python
# Extract font specification from DOCX
from docx import Document

doc = Document('mugi_sdl2_progress.docx')

# Find and print section 3-2 and surrounding content
in_font_section = False
line_count = 0

for para in doc.paragraphs:
    text = para.text
    
    # Check for section markers
    if '3-2' in text or 'フォント' in text or '変換テーブル' in text:
        in_font_section = True
    
    if in_font_section:
        print(text)
        line_count += 1
        if line_count > 100:  # Limit output
            break

# Also check tables
print("\n--- TABLES ---")
for table in doc.tables:
    for row in table.rows:
        print(" | ".join([cell.text for cell in row.cells]))
